from fastapi import APIRouter, File, UploadFile, HTTPException
from fastapi.responses import FileResponse, StreamingResponse
from docx import Document
from utils.render_imgs import render_image, render_image_full
from io import BytesIO

WORD_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

route = APIRouter(prefix="/img")

@route.post('/')
async def img(img : UploadFile = File(...)):
    resized_img = await render_image_full(img)
    resized_img.seek(0)
    return StreamingResponse(resized_img, media_type='image/jpeg')

# @route.post('/docx')
# async def docs(docx : UploadFile = File(...)):
#     if docx.content_type != WORD_CONTENT_TYPE:
#         raise HTTPException(422)
#     document = Document(BytesIO(await docx.read()))
#     for p in document.paragraphs:
#         print(p.text)
#         print(p)

#     for shape in document.inline_shapes:
#         print(shape.width)
#         print(shape.height)
#         print(shape.hyperlink)

#     return docx.filename

# docx.sections
# docx.paragraphs
# docx.inline_shapes
# docx.paragraphs._p